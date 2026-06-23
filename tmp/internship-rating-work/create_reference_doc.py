# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/luan/apps/ppm/tmp/internship-rating-work/Edijs_Senbergs_Internship_Supervisor_Reference.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "4B5563"
BORDER = "D7DBE2"
HEADER_FILL = "F2F4F7"
LIGHT_FILL = "F8FAFC"
SELECTED_FILL = "E8F2FF"


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_text(paragraph, text, size=11, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    if text:
        add_text(p, text, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(p, before=10, after=5)
        size, color = 15, BLUE
    elif level == 2:
        set_paragraph_spacing(p, before=12, after=6)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(p, before=8, after=4)
        size, color = 12, DARK_BLUE
    add_text(p, text, size=size, color=color, bold=True)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = width_twips


def set_table_geometry(table, widths_twips, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_twips[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_row_min_height(row, height_twips=360):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")


def cell_text(cell, text, size=10, color=INK, bold=False, italic=False, align=None):
    p = cell.paragraphs[0]
    p.text = ""
    set_paragraph_spacing(p, before=2, after=2, line=1.08)
    if align is not None:
        p.alignment = align
    add_text(p, text, size=size, color=color, bold=bold, italic=italic)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_metadata_table(doc):
    rows = [
        ("Company data (particulars from the Agreement)", "As per Internship Agreement"),
        ("Student", "Edijs Šenbergs"),
        ("Internship period", "As per Internship Agreement"),
        ("ECTS", "As per Internship Agreement"),
        (
            "Study program",
            'Joint Professional Bachelor study program "Finance Management Information Systems" of University of Latvia and Riga Technical University',
        ),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2600, 6760])
    for idx, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[idx].cells
        set_cell_shading(label_cell, HEADER_FILL)
        set_cell_shading(value_cell, "FFFFFF")
        cell_text(label_cell, label, size=9.5, color=MUTED, bold=True)
        cell_text(value_cell, value, size=10.2, color=INK)
        set_row_cant_split(table.rows[idx])
        set_row_min_height(table.rows[idx], 420)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_rating_table(doc, rows):
    headers = ["Assessment area", "Low", "Average", "Good", "Excellent", "N/A"]
    widths = [5100, 700, 900, 800, 1120, 740]
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_header(header)
    set_row_min_height(header, 360)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        cell_text(
            cell,
            label,
            size=9.2 if idx == 0 else 8.5,
            color=INK,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER,
        )

    for r_idx, (area, selected) in enumerate(rows, start=1):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        set_row_min_height(table.rows[-1], 360)
        body_fill = "FFFFFF" if r_idx % 2 else LIGHT_FILL
        for c_idx, cell in enumerate(cells):
            set_cell_shading(cell, body_fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_text(cells[0], area, size=9.0, color=INK)
        for c_idx, option in enumerate(headers[1:], start=1):
            mark = "X" if option == selected else ""
            if mark:
                set_cell_shading(cells[c_idx], SELECTED_FILL)
            cell_text(
                cells[c_idx],
                mark,
                size=10.0,
                color=DARK_BLUE if mark else MUTED,
                bold=bool(mark),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_score_table(doc):
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [6380, 2980])
    set_cell_shading(table.rows[0].cells[0], HEADER_FILL)
    set_cell_shading(table.rows[0].cells[1], HEADER_FILL)
    cell_text(table.rows[0].cells[0], "Assessment of Student's performance according to 10 (ten)-point grading scale", size=10, bold=True)
    cell_text(table.rows[0].cells[1], "Final assessment", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    scale = "10 - outstanding; 9 - excellent; 8 - very good; 7 - good; 6 - almost good; 5 - average; 4 - almost average; -1 - negative assessment"
    cell_text(table.rows[1].cells[0], scale, size=9.3, color=MUTED)
    cell_text(table.rows[1].cells[1], "8.5 / 10", size=16, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_row_cant_split(table.rows[0])
    set_row_cant_split(table.rows[1])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_signature_table(doc):
    rows = [
        ("Date", "19 June 2026"),
        ("Internship Supervisor's data from the Agreement / Signature", "Luan Pham"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [3400, 5960])
    for idx, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[idx].cells
        set_cell_shading(label_cell, HEADER_FILL if idx in (0, 2) else LIGHT_FILL)
        set_cell_shading(value_cell, "FFFFFF")
        cell_text(label_cell, label, size=9.5, color=MUTED, bold=True)
        cell_text(value_cell, value, size=10.5, color=INK, bold=("Signature" in label))
        set_row_cant_split(table.rows[idx])


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, after=0)
    add_text(header, "Internship Supervisor Reference | Edijs Šenbergs", size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, after=0)
    add_text(footer, "Page ", size=9, color=MUTED)
    add_page_field(footer)

    add_para(doc, "Internship Supervisor Reference", size=10.5, color=BLUE, bold=True, after=2)
    add_para(doc, "Internship Supervisor Reference Form", size=20, color=INK, bold=True, after=3)
    add_para(
        doc,
        "Student: Edijs Šenbergs | Supervisor: Luan Pham | Assessment: 8.5 / 10",
        size=11.2,
        color=MUTED,
        after=10,
    )

    add_metadata_table(doc)

    intro = (
        'On the Internship (data from the Agreement - in the period from As per Internship Agreement '
        'to As per Internship Agreement, ECTS: As per Internship Agreement) of the Student Edijs Šenbergs '
        'of Joint Professional Bachelor study program "Finance Management Information Systems" of University '
        'of Latvia and Riga Technical University (particulars from the Agreement - the study program).'
    )
    add_para(doc, intro, size=9.8, color=INK, after=6)

    add_heading(doc, "Skills in the respective professional field:", level=1)
    add_rating_table(
        doc,
        [
            ("Theoretical background, understanding of the main regularities of the professional field of activity", "Good"),
            ("Practical background, competences and skills applied during the Internship completing the assigned tasks", "Good"),
            ("Competence completing the tasks of the Internship", "Excellent"),
            ("Professional and proper performance, high quality work performance", "Good"),
            ("Ability to analyze, structure and process professional information", "Good"),
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Personal attitude and contribution of the Student:", level=1)
    add_rating_table(
        doc,
        [
            ("Responsible attitude towards the assigned duties, discipline", "Excellent"),
            ("Analytical, logical actions and situation understanding completing assigned tasks", "Good"),
            ("Autonomy and critical assessment of own performance", "Good"),
            ("Personal initiative and innovative approach", "Excellent"),
            ("Communication skills and ability to conduct argumentative discussion", "Excellent"),
            ("Ability to integrate and work in team", "Excellent"),
            ("Professional ethics", "Excellent"),
        ],
    )

    add_heading(doc, "Written reference in free form", level=1)
    add_para(
        doc,
        "Written reference in free form (including what materials the Student got acquainted with and what tools, equipment and software the Student used during the Internship, what projects or research the Student took part in, etc.).",
        size=9.5,
        color=MUTED,
        italic=True,
        after=8,
    )
    add_para(
        doc,
        "During the internship, Edijs Šenbergs got familiar with our internal workflow, task documentation, coordination process, and the tools used by the team. He was willing to take on assigned tasks and followed guidance carefully. His technical knowledge is still developing, but he showed a good learning attitude, asked questions when needed, and improved through feedback and practice. Edijs communicated well with colleagues, worked well with the team, and was also willing to help guide other interns. Overall, his practical skills are solid for this internship stage, with room to keep improving his technical depth, consistency, and independent problem solving.",
        size=10.5,
        color=INK,
        after=8,
    )

    add_heading(doc, "Overall assessment", level=1)
    add_score_table(doc)

    add_heading(doc, "Date and supervisor signature", level=1)
    add_signature_table(doc)

    doc.core_properties.author = "Luan Pham"
    doc.core_properties.title = "Internship Supervisor Reference Form - Edijs Senbergs"
    doc.core_properties.subject = "Internship Supervisor Reference"
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
